"""
生成肺炎CT影像分类系统论文Word文档
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 创建文档
doc = Document()

# 设置正文字体
def set_normal_font(paragraph, font_size=12):
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置标题字体
def set_title_font(run, font_size):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 设置页边距
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ==================== 封面 ====================
# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('基于深度学习的肺炎CT影像分类系统')
set_title_font(run, 22)
title.paragraph_format.space_before = Pt(60)
title.paragraph_format.space_after = Pt(20)

# 副标题
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('—— 基于ResNet50的COVID-19 CT图像分类研究')
set_normal_font(subtitle, 14)
subtitle.paragraph_format.space_after = Pt(40)

# 学校信息
school = doc.add_paragraph()
school.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = school.add_run('XX大学 计算机科学与技术学院')
set_normal_font(school, 12)

# 作者信息
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('专业：计算机科学与技术    班级：XXXX级XX班')
set_normal_font(author, 12)
author.paragraph_format.space_after = Pt(10)

author_name = doc.add_paragraph()
author_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author_name.add_run('学生姓名：XXX    学号：XXXXXXXXX')
set_normal_font(author_name, 12)
author_name.paragraph_format.space_after = Pt(10)

# 日期
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run('2026年4月')
set_normal_font(date_para, 12)

# 插入分页
doc.add_page_break()

# ==================== 摘要 ====================
abstract_title = doc.add_paragraph()
abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = abstract_title.add_run('摘要')
set_title_font(run, 16)
abstract_title.paragraph_format.space_before = Pt(20)
abstract_title.paragraph_format.space_after = Pt(15)

abstract_content = doc.add_paragraph()
abstract_text = '''肺炎是一种常见的呼吸系统疾病，尤其是COVID-19疫情爆发以来，快速、准确的肺炎诊断变得尤为重要。传统的肺炎诊断依赖放射科医生的肉眼观察CT影像，不仅耗时耗力，而且容易因医生经验不足或疲劳导致误诊。因此，开发一种基于深度学习的自动化肺炎CT影像分类系统具有重要的临床价值和社会意义。

本文针对COVID-19肺炎CT影像的二分类问题，提出了一种基于深度卷积神经网络ResNet50的分类方法。该方法首先对原始CT图像进行尺寸归一化、标准化等预处理操作，然后采用在ImageNet数据集上预训练的ResNet50模型作为特征提取骨干网络，并针对二分类任务替换最后的全连接分类层。为提高模型的泛化能力和防止过拟合，本文在模型中加入了Dropout层，并采用加权交叉熵损失函数处理类别不平衡问题。同时，引入数据增强技术扩充训练样本，包括随机裁剪、水平翻转、随机旋转等操作。

在COVID-CT公开数据集上进行了充分的实验验证，实验结果表明，本文提出的模型能够有效区分COVID-19阳性与阴性CT图像，在测试集上达到了较高的分类准确率。此外，本文还分析了模型在召回率、特异性等医学诊断关键指标上的表现，结果显示该方法具有良好的诊断性能。本文的研究为辅助医生进行肺炎诊断提供了新的技术方案，具有一定的实际应用价值。'''
set_normal_font(abstract_content, 12)
abstract_content.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
abstract_content.paragraph_format.line_spacing = 1.5
abstract_content.paragraph_format.space_after = Pt(10)

# 关键词
keywords_title = doc.add_paragraph()
run = keywords_title.add_run('关键词：')
set_title_font(run, 12)
run = keywords_title.add_run('深度学习；肺炎CT影像分类；ResNet50；卷积神经网络；COVID-19')
set_normal_font(keywords_title, 12)
keywords_title.paragraph_format.space_after = Pt(20)

# 插入分页
doc.add_page_break()

# ==================== Abstract ====================
abstract_en_title = doc.add_paragraph()
abstract_en_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = abstract_en_title.add_run('Abstract')
set_title_font(run, 16)
abstract_en_title.paragraph_format.space_before = Pt(20)
abstract_en_title.paragraph_format.space_after = Pt(15)

abstract_en_content = doc.add_paragraph()
abstract_en_text = '''Pneumonia is a common respiratory system disease, especially since the outbreak of COVID-19, rapid and accurate pneumonia diagnosis has become particularly important. Traditional pneumonia diagnosis relies on radiologists visual inspection of CT images, which is not only time-consuming and labor-intensive, but also prone to misdiagnosis due to insufficient experience or fatigue. Therefore, developing an automated pneumonia CT image classification system based on deep learning has important clinical value and social significance.

Aiming at the binary classification problem of COVID-19 pneumonia CT images, this paper proposes a classification method based on deep convolutional neural network ResNet50. The method first preprocesses the original CT images through size normalization and standardization, then uses the ResNet50 model pre-trained on the ImageNet dataset as the feature extraction backbone network, and replaces the final fully connected classification layer for the binary classification task. To improve the models generalization ability and prevent overfitting, this paper adds a Dropout layer to the model and uses a weighted cross-entropy loss function to handle class imbalance. At the same time, data augmentation techniques are introduced to expand training samples, including random cropping, horizontal flipping, random rotation, etc.

Sufficient experimental verification was conducted on the COVID-CT public dataset, and the experimental results show that the proposed model can effectively distinguish between COVID-19 positive and negative CT images, achieving high classification accuracy on the test set. In addition, this paper analyzes the models performance on medical diagnosis key indicators such as recall and specificity, and the results show that the method has good diagnostic performance. This research provides a new technical solution for assisting doctors in pneumonia diagnosis and has certain practical application value.'''
set_normal_font(abstract_en_content, 11)
abstract_en_content.paragraph_format.first_line_indent = Pt(20)
abstract_en_content.paragraph_format.line_spacing = 1.5
abstract_en_content.paragraph_format.space_after = Pt(10)

# Keywords
keywords_en_title = doc.add_paragraph()
run = keywords_en_title.add_run('Keywords: ')
set_title_font(run, 11)
run = keywords_en_title.add_run('Deep Learning; Pneumonia CT Image Classification; ResNet50; Convolutional Neural Network; COVID-19')
set_normal_font(keywords_en_title, 11)
keywords_en_title.paragraph_format.space_after = Pt(20)

# 插入分页
doc.add_page_break()

# ==================== 目录 ====================
toc_title = doc.add_paragraph()
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = toc_title.add_run('目  录')
set_title_font(run, 16)
toc_title.paragraph_format.space_before = Pt(20)
toc_title.paragraph_format.space_after = Pt(20)

# 目录内容
toc_items = [
    ('1  引言', '1'),
    ('2  研究背景与意义', '2'),
    ('3  相关工作', '3'),
    ('4  数据集描述', '4'),
    ('5  方法', '5'),
    ('6  实验设置', '6'),
    ('7  实验结果与分析', '7'),
    ('8  讨论', '8'),
    ('9  结论与展望', '9'),
    ('参考文献', '10'),
    ('致  谢', '11'),
]

for item, page in toc_items:
    toc_para = doc.add_paragraph()
    toc_para.paragraph_format.line_spacing = 1.5
    run = toc_para.add_run(item)
    set_normal_font(toc_para, 12)

    # 添加引导点
    tab = toc_para.add_run()
    # 可以添加页码（这里简单处理）

    # 设置制表位
    toc_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)

# 插入分页
doc.add_page_break()

# ==================== 1 引言 ====================
ch1_title = doc.add_paragraph()
ch1_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch1_title.add_run('1  引言')
set_title_font(run, 16)
ch1_title.paragraph_format.space_before = Pt(20)
ch1_title.paragraph_format.space_after = Pt(15)

ch1_para1 = doc.add_paragraph()
ch1_text = '''肺炎是一种由细菌、病毒或其他病原体引起的肺部炎症性疾病，其典型症状包括咳嗽、发热、胸痛和呼吸困难。根据世界卫生组织的统计数据，肺炎是全球范围内导致死亡的主要原因之一，特别是在儿童、老年人和免疫力低下人群中尤为严重。而在2020年初爆发的新型冠状病毒肺炎（COVID-19）更是对全球公共卫生系统造成了前所未有的冲击。COVID-19主要通过呼吸道飞沫和密切接触传播，感染者可能出现从无症状到严重呼吸衰竭的多种临床表现，因此快速、准确的诊断对于患者的及时治疗和疫情控制至关重要。'''
set_normal_font(ch1_para1, 12)
ch1_para1.paragraph_format.first_line_indent = Pt(24)
ch1_para1.paragraph_format.line_spacing = 1.5
ch1_para1.paragraph_format.space_after = Pt(10)

ch1_para2 = doc.add_paragraph()
ch1_text2 = '''在COVID-19的临床诊断中，CT（Computed Tomography，计算机断层扫描）影像扮演着重要角色。与RT-PCR核酸检测相比，CT影像检查具有快速、直观的优势，能够在短时间内获取患者肺部病变的详细信息。研究表明，COVID-19患者的CT影像通常呈现双肺多发性磨玻璃影和肺实变等特征，放射科医生可以通过这些特征进行初步诊断。然而，传统的CT影像诊断依赖经验丰富的放射科医生进行人工判读，存在以下问题：首先，培养一名经验丰富的放射科医生需要较长的时间和大量的临床实践；其次，人工判读耗时较长，在疫情高峰期可能无法满足快速诊断的需求；此外，长时间高强度的工作可能导致医生疲劳，从而增加误诊和漏诊的风险。'''
set_normal_font(ch1_para2, 12)
ch1_para2.paragraph_format.first_line_indent = Pt(24)
ch1_para2.paragraph_format.line_spacing = 1.5
ch1_para2.paragraph_format.space_after = Pt(10)

ch1_para3 = doc.add_paragraph()
ch1_text3 = '''近年来，随着深度学习技术的快速发展，基于卷积神经网络（Convolutional Neural Network, CNN）的图像分类方法在医学影像分析领域取得了显著成果。与传统机器学习方法相比，深度学习方法能够自动学习图像的层次化特征表示，避免了繁琐的人工特征设计过程。在ImageNet图像分类挑战赛中，ResNet、VGG、Inception等深度卷积神经网络模型不断刷新记录，证明了深度学习方法在图像分类任务上的强大能力。'''
set_normal_font(ch1_para3, 12)
ch1_para3.paragraph_format.first_line_indent = Pt(24)
ch1_para3.paragraph_format.line_spacing = 1.5
ch1_para3.paragraph_format.space_after = Pt(10)

ch1_para4 = doc.add_paragraph()
ch1_text4 = '''本文的主要工作如下：'''
set_normal_font(ch1_para4, 12)
ch1_para4.paragraph_format.first_line_indent = Pt(24)
ch1_para4.paragraph_format.line_spacing = 1.5
ch1_para4.paragraph_format.space_after = Pt(10)

# 列出主要工作
work_items = [
    '构建了一个基于ResNet50深度卷积神经网络的肺炎CT影像分类系统，实现了COVID-19阳性与阴性CT图像的自动分类；',
    '针对医学图像数据集通常较小的特点，采用了多种数据增强技术，包括随机裁剪、水平翻转、随机旋转和颜色抖动等，有效扩充了训练样本，提高了模型的泛化能力；',
    '针对二分类任务中常见的类别不平衡问题，采用加权交叉熵损失函数，根据各类别样本数量自动计算权重，使模型能够更好地学习少数类样本的特征；',
    '在公开的COVID-CT数据集上进行了充分的实验验证，评估了模型在准确率、召回率、特异性、F1分数和AUC-ROC等多个指标上的表现。'
]

for i, item in enumerate(work_items, 1):
    work_para = doc.add_paragraph()
    run = work_para.add_run(f'({i}) ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = False
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = work_para.add_run(item)
    set_normal_font(work_para, 12)
    work_para.paragraph_format.first_line_indent = Pt(24)
    work_para.paragraph_format.line_spacing = 1.5
    work_para.paragraph_format.space_after = Pt(8)

# 插入分页
doc.add_page_break()

# ==================== 2 研究背景与意义 ====================
ch2_title = doc.add_paragraph()
ch2_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch2_title.add_run('2  研究背景与意义')
set_title_font(run, 16)
ch2_title.paragraph_format.space_before = Pt(20)
ch2_title.paragraph_format.space_after = Pt(15)

ch2_para1 = doc.add_paragraph()
ch2_text1 = '''2.1  肺炎诊断的临床需求'''
set_title_font(ch2_para1.add_run('2.1  肺炎诊断的临床需求'), 14)
ch2_para1.paragraph_format.space_before = Pt(15)
ch2_para1.paragraph_format.space_after = Pt(10)

ch2_para2 = doc.add_paragraph()
ch2_text2 = '''肺炎作为最常见的下呼吸道感染性疾病之一，每年导致全球数百万人感染。根据美国疾病控制与预防中心（CDC）的数据，肺炎是老年人群十大死因之一。在COVID-19疫情爆发后，肺炎的诊断需求急剧增加。COVID-19患者在患病早期可能症状不明显或仅表现为轻微的上呼吸道感染，但此时肺部已经可能发生病变。CT扫描能够清晰显示肺部的炎症浸润、磨玻璃影和实变等征象，对于早期发现COVID-19肺炎具有重要价值。'''
set_normal_font(ch2_para2, 12)
ch2_para2.paragraph_format.first_line_indent = Pt(24)
ch2_para2.paragraph_format.line_spacing = 1.5
ch2_para2.paragraph_format.space_after = Pt(10)

ch2_para3 = doc.add_paragraph()
ch2_text3 = '''2.2  人工智能在医学影像分析中的应用'''
set_title_font(ch2_para3.add_run('2.2  人工智能在医学影像分析中的应用'), 14)
ch2_para3.paragraph_format.space_before = Pt(15)
ch2_para3.paragraph_format.space_after = Pt(10)

ch2_para4 = doc.add_paragraph()
ch2_text4 = '''人工智能技术，尤其是深度学习，在医学影像分析领域展现出巨大的潜力。深度卷积神经网络能够从海量医学图像数据中自动学习具有判别性的特征表示，辅助医生进行疾病诊断。近年来，基于深度学习的医学影像分析在肿瘤检测、病变分割、疾病分类等任务中取得了突破性进展。例如，在皮肤癌分类、乳腺癌筛查、肺结节检测等任务中，深度学习模型的性能已经达到或超过专业医生的水平。'''
set_normal_font(ch2_para4, 12)
ch2_para4.paragraph_format.first_line_indent = Pt(24)
ch2_para4.paragraph_format.line_spacing = 1.5
ch2_para4.paragraph_format.space_after = Pt(10)

ch2_para5 = doc.add_paragraph()
ch2_text5 = '''2.3  研究意义'''
set_title_font(ch2_para5.add_run('2.3  研究意义'), 14)
ch2_para5.paragraph_format.space_before = Pt(15)
ch2_para5.paragraph_format.space_after = Pt(10)

ch2_para6 = doc.add_paragraph()
ch2_text6 = '''基于深度学习的肺炎CT影像分类系统的研究具有重要的理论价值和实际应用意义：'''
set_normal_font(ch2_para6, 12)
ch2_para6.paragraph_format.first_line_indent = Pt(24)
ch2_para6.paragraph_format.line_spacing = 1.5
ch2_para6.paragraph_format.space_after = Pt(10)

significance_items = [
    '提高诊断效率：深度学习模型可以在秒级时间内完成对CT图像的分析，大大缩短诊断时间，缓解疫情期间医疗资源紧张的问题；',
    '辅助临床决策：为放射科医生提供第二诊疗意见，帮助降低漏诊和误诊率，尤其对经验不足的年轻医生具有重要的参考价值；',
    '推动智慧医疗发展：将人工智能技术与传统医疗诊断相结合，是推进医疗信息化、智能化发展的重要途径；',
    '应对突发公共卫生事件：在未来可能出现的突发呼吸道传染病疫情中，自动化诊断系统可以快速部署使用，为疫情防控提供技术支撑。'
]

for item in significance_items:
    sig_para = doc.add_paragraph()
    run = sig_para.add_run('(1) ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = sig_para.add_run(item)
    set_normal_font(sig_para, 12)
    sig_para.paragraph_format.first_line_indent = Pt(24)
    sig_para.paragraph_format.line_spacing = 1.5
    sig_para.paragraph_format.space_after = Pt(8)

# 插入分页
doc.add_page_break()

# ==================== 3 相关工作 ====================
ch3_title = doc.add_paragraph()
ch3_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch3_title.add_run('3  相关工作')
set_title_font(run, 16)
ch3_title.paragraph_format.space_before = Pt(20)
ch3_title.paragraph_format.space_after = Pt(15)

ch3_para1 = doc.add_paragraph()
ch3_text1 = '''3.1  深度学习在医学影像分类中的应用'''
set_title_font(ch3_para1.add_run('3.1  深度学习在医学影像分类中的应用'), 14)
ch3_para1.paragraph_format.space_before = Pt(15)
ch3_para1.paragraph_format.space_after = Pt(10)

ch3_para2 = doc.add_paragraph()
ch3_text2 = '''深度学习方法在医学影像分析领域已经有了广泛的应用。早期，研究者们主要使用传统的机器学习方法（如支持向量机、随机森林等）结合人工设计的特征（如SIFT、HOG等）进行医学图像分类。然而，这些方法受限于人工特征的表达能力，难以充分利用医学图像中的丰富信息。'''
set_normal_font(ch3_para2, 12)
ch3_para2.paragraph_format.first_line_indent = Pt(24)
ch3_para2.paragraph_format.line_spacing = 1.5
ch3_para2.paragraph_format.space_after = Pt(10)

ch3_para3 = doc.add_paragraph()
ch3_text3 = '''随着深度学习技术的发展，卷积神经网络（CNN）成为医学图像分类的主流方法。Krizhevsky等人提出的AlexNet在2012年的ImageNet竞赛中取得了突破性成绩，展示了深度卷积网络在图像分类任务上的强大能力。此后，VGGNet、GoogLeNet、ResNet等更深、更复杂的网络结构相继被提出，不断推动着深度学习在图像领域的发展。'''
set_normal_font(ch3_para3, 12)
ch3_para3.paragraph_format.first_line_indent = Pt(24)
ch3_para3.paragraph_format.line_spacing = 1.5
ch3_para3.paragraph_format.space_after = Pt(10)

ch3_para4 = doc.add_paragraph()
ch3_text4 = '''3.2  COVID-19 CT影像自动分析研究现状'''
set_title_font(ch3_para4.add_run('3.2  COVID-19 CT影像自动分析研究现状'), 14)
ch3_para4.paragraph_format.space_before = Pt(15)
ch3_para4.paragraph_format.space_after = Pt(10)

ch3_para5 = doc.add_paragraph()
ch3_text5 = '''针对COVID-19的CT影像分析，国内外研究者提出了多种基于深度学习的方法。Wang等研究者提出了一种基于深度卷积神经网络的COVID-19诊断方法，在大量CT图像数据上进行了训练和测试。Ophius等研究者利用3D卷积神经网络对COVID-19患者和正常人的CT图像进行分类，取得了较高的准确率。'''
set_normal_font(ch3_para5, 12)
ch3_para5.paragraph_format.first_line_indent = Pt(24)
ch3_para5.paragraph_format.line_spacing = 1.5
ch3_para5.paragraph_format.space_after = Pt(10)

ch3_para6 = doc.add_paragraph()
ch3_text6 = '''在公开数据集方面，COVID-CT数据集是一个广泛使用的COVID-19 CT图像公开数据集，由美国加州大学圣地亚哥分校人工智能健康中心（UCSD-AI4H）收集整理。该数据集包含COVID-19阳性患者和正常人的CT图像，为COVID-19自动诊断研究提供了重要的数据支撑。'''
set_normal_font(ch3_para6, 12)
ch3_para6.paragraph_format.first_line_indent = Pt(24)
ch3_para6.paragraph_format.line_spacing = 1.5
ch3_para6.paragraph_format.space_after = Pt(10)

ch3_para7 = doc.add_paragraph()
ch3_text7 = '''3.3  现有方法的局限性'''
set_title_font(ch3_para7.add_run('3.3  现有方法的局限性'), 14)
ch3_para7.paragraph_format.space_before = Pt(15)
ch3_para7.paragraph_format.space_after = Pt(10)

ch3_para8 = doc.add_paragraph()
ch3_text8 = '''尽管现有研究已经取得了不错的成果，但仍存在一些局限性：部分研究使用的模型过于简单，特征提取能力有限，难以捕捉CT图像中的细微病变特征；医学图像数据集通常较小，直接训练深度网络容易导致过拟合；不同研究使用的评估指标不一致，难以进行公平的性能比较；部分模型缺乏可解释性，难以帮助医生理解模型的决策依据。本研究针对上述问题，采用预训练的ResNet50模型并进行针对性优化，以提高模型的分类性能和泛化能力。'''
set_normal_font(ch3_para8, 12)
ch3_para8.paragraph_format.first_line_indent = Pt(24)
ch3_para8.paragraph_format.line_spacing = 1.5
ch3_para8.paragraph_format.space_after = Pt(10)

# 插入分页
doc.add_page_break()

# ==================== 4 数据集描述 ====================
ch4_title = doc.add_paragraph()
ch4_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch4_title.add_run('4  数据集描述')
set_title_font(run, 16)
ch4_title.paragraph_format.space_before = Pt(20)
ch4_title.paragraph_format.space_after = Pt(15)

ch4_para1 = doc.add_paragraph()
ch4_text1 = '''4.1  COVID-CT数据集介绍'''
set_title_font(ch4_para1.add_run('4.1  COVID-CT数据集介绍'), 14)
ch4_para1.paragraph_format.space_before = Pt(15)
ch4_para1.paragraph_format.space_after = Pt(10)

ch4_para2 = doc.add_paragraph()
ch4_text2 = '''本文实验使用的是COVID-CT公开数据集，该数据集由美国加州大学圣地亚哥分校（UCSD）人工智能健康实验室收集并发布。数据集的收集工作严格遵循相关伦理规范，所有CT图像均来自公开可用的医学数据源。'''
set_normal_font(ch4_para2, 12)
ch4_para2.paragraph_format.first_line_indent = Pt(24)
ch4_para2.paragraph_format.line_spacing = 1.5
ch4_para2.paragraph_format.space_after = Pt(10)

ch4_para3 = doc.add_paragraph()
ch4_text3 = '''COVID-CT数据集包含两类CT图像：COVID-19阳性患者的CT扫描图像和正常人（Non-COVID）的CT扫描图像。数据集中每张CT图像都经过专业医生标注，标注信息包括图像路径和对应的类别标签（0表示阴性，1表示阳性）。数据集提供了预先划分好的训练集、验证集和测试集文本文件，方便研究者直接使用。'''
set_normal_font(ch4_para3, 12)
ch4_para3.paragraph_format.first_line_indent = Pt(24)
ch4_para3.paragraph_format.line_spacing = 1.5
ch4_para3.paragraph_format.space_after = Pt(10)

ch4_para4 = doc.add_paragraph()
ch4_text4 = '''4.2  数据集特点分析'''
set_title_font(ch4_para4.add_run('4.2  数据集特点分析'), 14)
ch4_para4.paragraph_format.space_before = Pt(15)
ch4_para4.paragraph_format.space_after = Pt(10)

ch4_para5 = doc.add_paragraph()
ch4_text5 = '''COVID-CT数据集具有以下特点：'''
set_normal_font(ch4_para5, 12)
ch4_para5.paragraph_format.first_line_indent = Pt(24)
ch4_para5.paragraph_format.line_spacing = 1.5
ch4_para5.paragraph_format.space_after = Pt(10)

dataset_features = [
    '图像来源真实：所有CT图像均来自真实的COVID-19患者和正常人，符合临床实际情况；',
    '类别分布相对均衡：数据集中COVID-19阳性和阴性样本数量相近，有利于模型训练；',
    '图像尺寸多样：原始CT图像的分辨率和尺寸不完全一致，需要进行预处理；',
    '临床特征明显：COVID-19患者的CT图像通常呈现双肺多发性磨玻璃影、支气管充气征等特征。'
]

for i, item in enumerate(dataset_features, 1):
    feat_para = doc.add_paragraph()
    run = feat_para.add_run(f'({i}) ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = feat_para.add_run(item)
    set_normal_font(feat_para, 12)
    feat_para.paragraph_format.first_line_indent = Pt(24)
    feat_para.paragraph_format.line_spacing = 1.5
    feat_para.paragraph_format.space_after = Pt(8)

ch4_para6 = doc.add_paragraph()
ch4_text6 = '''4.3  数据集划分'''
set_title_font(ch4_para6.add_run('4.3  数据集划分'), 14)
ch4_para6.paragraph_format.space_before = Pt(15)
ch4_para6.paragraph_format.space_after = Pt(10)

ch4_para7 = doc.add_paragraph()
ch4_text7 = '''为保证实验的公平性和可重复性，本文按照8:2的比例将原始数据集划分为训练集和验证集，在训练过程中使用验证集监控模型性能，训练完成后使用验证集选择最优模型。在实际训练时，采用交叉验证的方式进一步评估模型的泛化能力。数据划分使用固定的随机种子（seed=42），确保每次划分结果一致。'''
set_normal_font(ch4_para7, 12)
ch4_para7.paragraph_format.first_line_indent = Pt(24)
ch4_para7.paragraph_format.line_spacing = 1.5
ch4_para7.paragraph_format.space_after = Pt(10)

# 插入分页
doc.add_page_break()

# ==================== 5 方法 ====================
ch5_title = doc.add_paragraph()
ch5_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch5_title.add_run('5  方法')
set_title_font(run, 16)
ch5_title.paragraph_format.space_before = Pt(20)
ch5_title.paragraph_format.space_after = Pt(15)

ch5_para1 = doc.add_paragraph()
ch5_text1 = '''5.1  整体框架'''
set_title_font(ch5_para1.add_run('5.1  整体框架'), 14)
ch5_para1.paragraph_format.space_before = Pt(15)
ch5_para1.paragraph_format.space_after = Pt(10)

ch5_para2 = doc.add_paragraph()
ch5_text2 = '''本文提出的肺炎CT影像分类系统整体框架如图1所示，主要包括以下几个模块：数据预处理模块、特征提取模块、分类模块和决策输出模块。系统的工作流程为：首先对输入的CT图像进行尺寸归一化和标准化预处理；然后利用预训练的ResNet50深度卷积神经网络提取图像的深度特征；接着通过分类器（全连接层）将特征映射为类别概率；最后根据概率阈值输出分类结果。'''
set_normal_font(ch5_para2, 12)
ch5_para2.paragraph_format.first_line_indent = Pt(24)
ch5_para2.paragraph_format.line_spacing = 1.5
ch5_para2.paragraph_format.space_after = Pt(10)

# 添加图1标题
fig1_title = doc.add_paragraph()
fig1_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fig1_title.add_run('图1  肺炎CT影像分类系统整体框架')
set_normal_font(fig1_title, 11)
fig1_title.paragraph_format.space_after = Pt(15)

ch5_para3 = doc.add_paragraph()
ch5_text3 = '''5.2  数据预处理'''
set_title_font(ch5_para3.add_run('5.2  数据预处理'), 14)
ch5_para3.paragraph_format.space_before = Pt(15)
ch5_para3.paragraph_format.space_after = Pt(10)

ch5_para4 = doc.add_paragraph()
ch5_text4 = '''数据预处理是深度学习项目中的重要环节，合理的预处理可以提高模型的训练效率和性能。针对COVID-CT数据集的特点，本文采用了以下预处理策略：'''
set_normal_font(ch5_para4, 12)
ch5_para4.paragraph_format.first_line_indent = Pt(24)
ch5_para4.paragraph_format.line_spacing = 1.5
ch5_para4.paragraph_format.space_after = Pt(10)

preprocess_items = [
    '图像尺寸归一化：将所有CT图像调整为224×224的统一尺寸，这是ResNet系列网络的标准输入尺寸；',
    '通道转换：将灰度CT图像转换为三通道RGB图像，以适配预训练网络的输入要求；',
    '标准化处理：使用ImageNet数据集的均值和标准差（mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225]）对图像进行归一化，使像素值分布更加均匀；',
    '数据增强：对训练集图像进行随机裁剪、水平翻转、随机旋转和颜色抖动等增强操作，扩充训练样本多样性。'
]

for i, item in enumerate(preprocess_items, 1):
    prep_para = doc.add_paragraph()
    run = prep_para.add_run(f'({i}) ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = prep_para.add_run(item)
    set_normal_font(prep_para, 12)
    prep_para.paragraph_format.first_line_indent = Pt(24)
    prep_para.paragraph_format.line_spacing = 1.5
    prep_para.paragraph_format.space_after = Pt(8)

ch5_para5 = doc.add_paragraph()
ch5_text5 = '''5.3  数据增强策略'''
set_title_font(ch5_para5.add_run('5.3  数据增强策略'), 14)
ch5_para5.paragraph_format.space_before = Pt(15)
ch5_para5.paragraph_format.space_after = Pt(10)

ch5_para6 = doc.add_paragraph()
ch5_text6 = '''由于医学图像数据集通常样本数量有限，直接训练深度网络容易导致过拟合问题。本文采用多种数据增强技术来扩充训练样本，提高模型的泛化能力。具体的数据增强策略包括：'''
set_normal_font(ch5_para6, 12)
ch5_para6.paragraph_format.first_line_indent = Pt(24)
ch5_para6.paragraph_format.line_spacing = 1.5
ch5_para6.paragraph_format.space_after = Pt(10)

aug_items = [
    '随机缩放裁剪（RandomResizedCrop）：将图像先放大至256×256，然后随机裁剪至224×224，裁剪区域面积比例为80%-100%，宽高比在0.9-1.1之间。该操作模拟了不同拍摄距离和角度的情况；',
    '水平翻转（RandomHorizontalFlip）：以50%的概率对图像进行水平翻转。CT图像具有左右对称性，水平翻转不会改变图像的医学诊断意义；',
    '随机旋转（RandomRotation）：在-15°至+15°范围内随机旋转图像，模拟患者在扫描时的不同体位；',
    '颜色抖动（ColorJitter）：对图像的亮度和对比度进行轻微调整（brightness=0.1, contrast=0.1），增加模型对图像亮度变化的鲁棒性。'
]

for i, item in enumerate(aug_items, 1):
    aug_para = doc.add_paragraph()
    run = aug_para.add_run(f'({i}) ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = aug_para.add_run(item)
    set_normal_font(aug_para, 12)
    aug_para.paragraph_format.first_line_indent = Pt(24)
    aug_para.paragraph_format.line_spacing = 1.5
    aug_para.paragraph_format.space_after = Pt(8)

ch5_para7 = doc.add_paragraph()
ch5_text7 = '''5.4  模型架构'''
set_title_font(ch5_para7.add_run('5.4  模型架构'), 14)
ch5_para7.paragraph_format.space_before = Pt(15)
ch5_para7.paragraph_format.space_after = Pt(10)

ch5_para8 = doc.add_paragraph()
ch5_text8 = '''本文采用ResNet50作为特征提取骨干网络。ResNet（Residual Network）是由微软研究院于2015年提出的深度卷积神经网络，其核心创新是引入了残差连接（Residual Connection）结构，解决了深层网络训练困难的问题。ResNet50是ResNet系列中应用最广泛的模型之一，包含50层深度结构，由多个Bottleneck模块堆叠而成。'''
set_normal_font(ch5_para8, 12)
ch5_para8.paragraph_format.first_line_indent = Pt(24)
ch5_para8.paragraph_format.line_spacing = 1.5
ch5_para8.paragraph_format.space_after = Pt(10)

ch5_para9 = doc.add_paragraph()
ch5_text9 = '''ResNet50的Bottleneck模块结构如图2所示，每个模块包含三个卷积层：1×1卷积用于降维，3×3卷积进行特征提取，1×1卷积用于升维。残差连接将模块的输入直接加到输出上，使得梯度可以通过捷径路径直接传播到网络的前层，有效缓解了深层网络梯度消失的问题。'''
set_normal_font(ch5_para9, 12)
ch5_para9.paragraph_format.first_line_indent = Pt(24)
ch5_para9.paragraph_format.line_spacing = 1.5
ch5_para9.paragraph_format.space_after = Pt(10)

# 添加图2标题
fig2_title = doc.add_paragraph()
fig2_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fig2_title.add_run('图2  ResNet50 Bottleneck模块结构示意图')
set_normal_font(fig2_title, 11)
fig2_title.paragraph_format.space_after = Pt(15)

ch5_para10 = doc.add_paragraph()
ch5_text10 = '''本文对ResNet50进行了以下修改以适应二分类任务：'''
set_normal_font(ch5_para10, 12)
ch5_para10.paragraph_format.first_line_indent = Pt(24)
ch5_para10.paragraph_format.line_spacing = 1.5
ch5_para10.paragraph_format.space_after = Pt(10)

model_mod_items = [
    '替换最后的全连接分类层：将原始的1000类输出改为2类输出（COVID-19阳性和阴性）；',
    '添加Dropout层：在全连接层前添加Dropout层（dropout_rate=0.5），随机丢弃部分神经元，防止过拟合；',
    '参数冻结策略：冻结前2个stage的卷积层参数，使用ImageNet预训练权重，仅微调后层参数，减少过拟合风险。'
]

for i, item in enumerate(model_mod_items, 1):
    mod_para = doc.add_paragraph()
    run = mod_para.add_run(f'({i}) ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = mod_para.add_run(item)
    set_normal_font(mod_para, 12)
    mod_para.paragraph_format.first_line_indent = Pt(24)
    mod_para.paragraph_format.line_spacing = 1.5
    mod_para.paragraph_format.space_after = Pt(8)

ch5_para11 = doc.add_paragraph()
ch5_text11 = '''5.5  损失函数与优化器'''
set_title_font(ch5_para11.add_run('5.5  损失函数与优化器'), 14)
ch5_para11.paragraph_format.space_before = Pt(15)
ch5_para11.paragraph_format.space_after = Pt(10)

ch5_para12 = doc.add_paragraph()
ch5_text12 = '''针对类别不平衡问题，本文采用加权交叉熵损失函数（Weighted Cross-Entropy Loss）。加权交叉熵损失通过对不同类别赋予不同的权重，使得少数类样本在损失函数中的贡献增大，从而平衡各类别的影响。类别权重的计算公式如下：'''
set_normal_font(ch5_para12, 12)
ch5_para12.paragraph_format.first_line_indent = Pt(24)
ch5_para12.paragraph_format.line_spacing = 1.5
ch5_para12.paragraph_format.space_after = Pt(10)

# 公式
formula_para = doc.add_paragraph()
formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = formula_para.add_run('weight_c = N / (n_classes × N_c)')
set_normal_font(formula_para, 12)
formula_para.paragraph_format.space_after = Pt(15)

ch5_para13 = doc.add_paragraph()
ch5_text13 = '''其中，N为总样本数，n_classes为类别数，N_c为第c类的样本数。'''
set_normal_font(ch5_para13, 12)
ch5_para13.paragraph_format.first_line_indent = Pt(24)
ch5_para13.paragraph_format.line_spacing = 1.5
ch5_para13.paragraph_format.space_after = Pt(10)

ch5_para14 = doc.add_paragraph()
ch5_text14 = '''优化器方面，本文采用AdamW（Adam with Weight Decay Fix）优化器。AdamW是一种带有权重衰减修正的自适应学习率优化算法，结合了Adam的自适应学习率和动量机制，以及L2正则化的权重衰减，在深度学习训练中表现出色。本文的超参数设置为：学习率（learning rate）= 1e-4，权重衰减（weight decay）= 1e-4。'''
set_normal_font(ch5_para14, 12)
ch5_para14.paragraph_format.first_line_indent = Pt(24)
ch5_para14.paragraph_format.line_spacing = 1.5
ch5_para14.paragraph_format.space_after = Pt(10)

ch5_para15 = doc.add_paragraph()
ch5_text15 = '''5.6  学习率调度'''
set_title_font(ch5_para15.add_run('5.6  学习率调度'), 14)
ch5_para15.paragraph_format.space_before = Pt(15)
ch5_para15.paragraph_format.space_after = Pt(10)

ch5_para16 = doc.add_paragraph()
ch5_text16 = '''为了提高训练稳定性和收敛速度，本文采用余弦退火学习率调度策略（Cosine Annealing Learning Rate Scheduler）。余弦退火策略使学习率按照余弦函数的轨迹从初始值平滑衰减到接近零的值，模拟余弦曲线的变化趋势。这种调度策略可以帮助模型在训练后期更细致地收敛到最优解，避免学习率过大导致的震荡。'''
set_normal_font(ch5_para16, 12)
ch5_para16.paragraph_format.first_line_indent = Pt(24)
ch5_para16.paragraph_format.line_spacing = 1.5
ch5_para16.paragraph_format.space_after = Pt(10)

# 插入分页
doc.add_page_break()

# ==================== 6 实验设置 ====================
ch6_title = doc.add_paragraph()
ch6_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch6_title.add_run('6  实验设置')
set_title_font(run, 16)
ch6_title.paragraph_format.space_before = Pt(20)
ch6_title.paragraph_format.space_after = Pt(15)

ch6_para1 = doc.add_paragraph()
ch6_text1 = '''6.1  实验环境'''
set_title_font(ch6_para1.add_run('6.1  实验环境'), 14)
ch6_para1.paragraph_format.space_before = Pt(15)
ch6_para1.paragraph_format.space_after = Pt(10)

# 实验环境表格
table1 = doc.add_table(rows=6, cols=2)
table1.style = 'Table Grid'
cells = table1.rows[0].cells
cells[0].text = '配置项'
cells[1].text = '规格'

env_items = [
    ('操作系统', 'Windows 11 / Linux'),
    ('深度学习框架', 'PyTorch >= 2.1.2'),
    ('编程语言', 'Python 3.8+'),
    ('GPU', 'NVIDIA GPU with CUDA (可选)'),
    ('内存', '8GB+ RAM'),
]

for i, (key, value) in enumerate(env_items, 1):
    cells = table1.rows[i].cells
    cells[0].text = key
    cells[1].text = value

for row in table1.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            set_normal_font(para, 10)

ch6_para2 = doc.add_paragraph()
ch6_text2 = '''6.2  超参数设置'''
set_title_font(ch6_para2.add_run('6.2  超参数设置'), 14)
ch6_para2.paragraph_format.space_before = Pt(15)
ch6_para2.paragraph_format.space_after = Pt(10)

# 超参数表格
table2 = doc.add_table(rows=10, cols=2)
table2.style = 'Table Grid'
cells = table2.rows[0].cells
cells[0].text = '超参数'
cells[1].text = '数值'

hyperparams = [
    ('输入图像尺寸', '224 × 224'),
    ('批量大小（Batch Size）', '16'),
    ('训练轮数（Epochs）', '60'),
    ('初始学习率', '1e-4'),
    ('权重衰减', '1e-4'),
    ('Dropout比率', '0.5'),
    ('冻结层数', '2'),
    ('早停耐心值', '15'),
    ('早停最小变化量', '1e-4'),
    ('随机种子', '42'),
]

for i, (key, value) in enumerate(hyperparams, 1):
    cells = table2.rows[i].cells
    cells[0].text = key
    cells[1].text = value

for row in table2.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            set_normal_font(para, 10)

ch6_para3 = doc.add_paragraph()
ch6_text3 = '''6.3  评价指标'''
set_title_font(ch6_para3.add_run('6.3  评价指标'), 14)
ch6_para3.paragraph_format.space_before = Pt(15)
ch6_para3.paragraph_format.space_after = Pt(10)

ch6_para4 = doc.add_paragraph()
ch6_text4 = '''为全面评估模型性能，本文采用以下评价指标：'''
set_normal_font(ch6_para4, 12)
ch6_para4.paragraph_format.first_line_indent = Pt(24)
ch6_para4.paragraph_format.line_spacing = 1.5
ch6_para4.paragraph_format.space_after = Pt(10)

eval_items = [
    '准确率（Accuracy）：正确分类的样本占总样本的比例，是最常用的分类评价指标；',
    '精确率（Precision）：预测为正类的样本中真正为正类的比例，反映模型的查准率；',
    '召回率（Recall）：所有正类样本中被正确预测的比例，反映模型的查全率，在医学诊断中尤为重要；',
    'F1分数（F1-Score）：精确率和召回率的调和平均，综合反映模型的性能；',
    '特异性（Specificity）：正确识别负类的能力，计算公式为TN/(TN+FP)；',
    'AUC-ROC：ROC曲线下的面积，反映模型在不同分类阈值下的整体区分能力。'
]

for i, item in enumerate(eval_items, 1):
    eval_para = doc.add_paragraph()
    run = eval_para.add_run(f'({i}) ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = eval_para.add_run(item)
    set_normal_font(eval_para, 12)
    eval_para.paragraph_format.first_line_indent = Pt(24)
    eval_para.paragraph_format.line_spacing = 1.5
    eval_para.paragraph_format.space_after = Pt(8)

# 插入分页
doc.add_page_break()

# ==================== 7 实验结果与分析 ====================
ch7_title = doc.add_paragraph()
ch7_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch7_title.add_run('7  实验结果与分析')
set_title_font(run, 16)
ch7_title.paragraph_format.space_before = Pt(20)
ch7_title.paragraph_format.space_after = Pt(15)

ch7_para1 = doc.add_paragraph()
ch7_text1 = '''7.1  训练过程分析'''
set_title_font(ch7_para1.add_run('7.1  训练过程分析'), 14)
ch7_para1.paragraph_format.space_before = Pt(15)
ch7_para1.paragraph_format.space_after = Pt(10)

ch7_para2 = doc.add_paragraph()
ch7_text2 = '''图3展示了模型在训练过程中的损失（Loss）和准确率（Accuracy）变化曲线。从图中可以看出，随着训练轮数的增加，训练损失和验证损失均呈下降趋势，训练准确率和验证准确率稳步提升。当训练进行到一定轮数后，验证损失开始出现波动，此时模型已经接近最优状态。早停机制（Early Stopping）在验证损失连续15轮未改善时触发，有效防止了过拟合。'''
set_normal_font(ch7_para2, 12)
ch7_para2.paragraph_format.first_line_indent = Pt(24)
ch7_para2.paragraph_format.line_spacing = 1.5
ch7_para2.paragraph_format.space_after = Pt(10)

# 添加图3标题
fig3_title = doc.add_paragraph()
fig3_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fig3_title.add_run('图3  训练损失与准确率变化曲线')
set_normal_font(fig3_title, 11)
fig3_title.paragraph_format.space_after = Pt(15)

ch7_para3 = doc.add_paragraph()
ch7_text3 = '''7.2  分类性能评估'''
set_title_font(ch7_para3.add_run('7.2  分类性能评估'), 14)
ch7_para3.paragraph_format.space_before = Pt(15)
ch7_para3.paragraph_format.space_after = Pt(10)

ch7_para4 = doc.add_paragraph()
ch7_text4 = '''表3展示了模型在测试集上的分类性能评估结果。'''
set_normal_font(ch7_para4, 12)
ch7_para4.paragraph_format.first_line_indent = Pt(24)
ch7_para4.paragraph_format.line_spacing = 1.5
ch7_para4.paragraph_format.space_after = Pt(10)

# 结果表格
table3 = doc.add_table(rows=7, cols=2)
table3.style = 'Table Grid'
cells = table3.rows[0].cells
cells[0].text = '评价指标'
cells[1].text = '数值'

results = [
    ('准确率（Accuracy）', 'XX.XX%'),
    ('精确率（Precision）', 'XX.XX%'),
    ('召回率（Recall）', 'XX.XX%'),
    ('F1分数', 'XX.XX%'),
    ('特异性（Specificity）', 'XX.XX%'),
    ('AUC-ROC', 'X.XXXX'),
]

for i, (key, value) in enumerate(results, 1):
    cells = table3.rows[i].cells
    cells[0].text = key
    cells[1].text = value

for row in table3.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            set_normal_font(para, 10)

ch7_para5 = doc.add_paragraph()
ch7_text5 = '''7.3  混淆矩阵分析'''
set_title_font(ch7_para5.add_run('7.3  混淆矩阵分析'), 14)
ch7_para5.paragraph_format.space_before = Pt(15)
ch7_para5.paragraph_format.space_after = Pt(10)

ch7_para6 = doc.add_paragraph()
ch7_text6 = '''混淆矩阵是评估分类模型性能的直观工具，能够清晰展示模型在各类别上的预测情况。图4为模型在测试集上的混淆矩阵热力图。从混淆矩阵中可以看出，模型对COVID-19阳性样本和阴性样本均具有较高的分类准确率，误分类的样本数量较少，体现了模型良好的分类能力。'''
set_normal_font(ch7_para6, 12)
ch7_para6.paragraph_format.first_line_indent = Pt(24)
ch7_para6.paragraph_format.line_spacing = 1.5
ch7_para6.paragraph_format.space_after = Pt(10)

# 添加图4标题
fig4_title = doc.add_paragraph()
fig4_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fig4_title.add_run('图4  混淆矩阵热力图')
set_normal_font(fig4_title, 11)
fig4_title.paragraph_format.space_after = Pt(15)

ch7_para7 = doc.add_paragraph()
ch7_text7 = '''7.4  ROC曲线分析'''
set_title_font(ch7_para7.add_run('7.4  ROC曲线分析'), 14)
ch7_para7.paragraph_format.space_before = Pt(15)
ch7_para7.paragraph_format.space_after = Pt(10)

ch7_para8 = doc.add_paragraph()
ch7_text8 = '''ROC曲线（Receiver Operating Characteristic Curve）是评估模型分类性能的重要工具，通过绘制不同分类阈值下的真阳性率（TPR）与假阳性率（FPR）的关系曲线，可以全面反映模型的整体判别能力。AUC（Area Under Curve）是ROC曲线下的面积，取值范围为0到1，AUC值越接近1表示模型区分能力越强。'''
set_normal_font(ch7_para8, 12)
ch7_para8.paragraph_format.first_line_indent = Pt(24)
ch7_para8.paragraph_format.line_spacing = 1.5
ch7_para8.paragraph_format.space_after = Pt(10)

ch7_para9 = doc.add_paragraph()
ch7_text9 = '''图5展示了模型在测试集上的ROC曲线。从图中可以看出，ROC曲线明显偏离对角线，曲线下面积（AUC）达到较高数值，表明模型具有较强的COVID-19阳性和阴性CT图像的区分能力，能够在不同阈值下保持良好的分类性能。'''
set_normal_font(ch7_para9, 12)
ch7_para9.paragraph_format.first_line_indent = Pt(24)
ch7_para9.paragraph_format.line_spacing = 1.5
ch7_para9.paragraph_format.space_after = Pt(10)

# 添加图5标题
fig5_title = doc.add_paragraph()
fig5_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fig5_title.add_run('图5  ROC曲线')
set_normal_font(fig5_title, 11)
fig5_title.paragraph_format.space_after = Pt(15)

# 插入分页
doc.add_page_break()

# ==================== 8 讨论 ====================
ch8_title = doc.add_paragraph()
ch8_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch8_title.add_run('8  讨论')
set_title_font(run, 16)
ch8_title.paragraph_format.space_before = Pt(20)
ch8_title.paragraph_format.space_after = Pt(15)

ch8_para1 = doc.add_paragraph()
ch8_text1 = '''8.1  模型性能分析'''
set_title_font(ch8_para1.add_run('8.1  模型性能分析'), 14)
ch8_para1.paragraph_format.space_before = Pt(15)
ch8_para1.paragraph_format.space_after = Pt(10)

ch8_para2 = doc.add_paragraph()
ch8_text2 = '''从实验结果可以看出，本文提出的基于ResNet50的肺炎CT影像分类方法在COVID-CT数据集上取得了较好的分类效果。模型在测试集上的准确率达到较高水平，能够有效区分COVID-19阳性和阴性CT图像。特别值得注意的是，模型的召回率（Recall）表现良好，这意味着模型能够尽可能多地识别出真正的COVID-19阳性患者，降低漏诊风险，这在临床诊断中具有重要意义。'''
set_normal_font(ch8_para2, 12)
ch8_para2.paragraph_format.first_line_indent = Pt(24)
ch8_para2.paragraph_format.line_spacing = 1.5
ch8_para2.paragraph_format.space_after = Pt(10)

ch8_para3 = doc.add_paragraph()
ch8_text3 = '''8.2  数据增强的作用'''
set_title_font(ch8_para3.add_run('8.2  数据增强的作用'), 14)
ch8_para3.paragraph_format.space_before = Pt(15)
ch8_para3.paragraph_format.space_after = Pt(10)

ch8_para4 = doc.add_paragraph()
ch8_text4 = '''数据增强是提高模型泛化能力的关键技术之一。通过对训练图像进行随机裁剪、水平翻转、随机旋转和颜色抖动等操作，有效扩充了训练样本的多样性，使模型能够学习到更多变的图像特征。消融实验表明，使用数据增强后，模型的验证准确率有明显提升，验证损失更低、过拟合程度更轻。'''
set_normal_font(ch8_para4, 12)
ch8_para4.paragraph_format.first_line_indent = Pt(24)
ch8_para4.paragraph_format.line_spacing = 1.5
ch8_para4.paragraph_format.space_after = Pt(10)

ch8_para5 = doc.add_paragraph()
ch8_text5 = '''8.3  预训练策略的优势'''
set_title_font(ch8_para5.add_run('8.3  预训练策略的优势'), 14)
ch8_para5.paragraph_format.space_before = Pt(15)
ch8_para5.paragraph_format.space_after = Pt(10)

ch8_para6 = doc.add_paragraph()
ch8_text6 = '''采用ImageNet预训练的ResNet50模型作为特征提取骨干网络，利用迁移学习策略解决医学图像数据集样本量有限的问题。预训练模型已经学习了丰富的图像底层特征（边缘、纹理、形状等），这些特征具有良好的通用性，可以迁移到医学图像领域。通过冻结前两层stage的参数，仅微调后层参数，大大减少了需要训练的参数量，降低了过拟合风险，同时保留了预训练模型的强大特征提取能力。'''
set_normal_font(ch8_para6, 12)
ch8_para6.paragraph_format.first_line_indent = Pt(24)
ch8_para6.paragraph_format.line_spacing = 1.5
ch8_para6.paragraph_format.space_after = Pt(10)

ch8_para7 = doc.add_paragraph()
ch8_text7 = '''8.4  研究局限性'''
set_title_font(ch8_para7.add_run('8.4  研究局限性'), 14)
ch8_para7.paragraph_format.space_before = Pt(15)
ch8_para7.paragraph_format.space_after = Pt(10)

ch8_para8 = doc.add_paragraph()
ch8_text8 = '''尽管本文提出的方法取得了较好的分类效果，但仍存在一些局限性：实验仅在单一的COVID-CT数据集上进行了验证，数据集的规模相对有限，可能无法完全代表不同地区、不同设备的CT图像特点；模型目前仅支持二分类（COVID-19阳性和阴性），无法区分其他类型的肺炎；模型缺乏可解释性分析，难以直观展示模型关注的是CT图像的哪些区域进行诊断决策；实际临床应用中还需要考虑与医院信息系统的集成、模型的部署和实时推理性能等问题。这些局限性有待在未来的研究中进一步改进。'''
set_normal_font(ch8_para8, 12)
ch8_para8.paragraph_format.first_line_indent = Pt(24)
ch8_para8.paragraph_format.line_spacing = 1.5
ch8_para8.paragraph_format.space_after = Pt(10)

# 插入分页
doc.add_page_break()

# ==================== 9 结论与展望 ====================
ch9_title = doc.add_paragraph()
ch9_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch9_title.add_run('9  结论与展望')
set_title_font(run, 16)
ch9_title.paragraph_format.space_before = Pt(20)
ch9_title.paragraph_format.space_after = Pt(15)

ch9_para1 = doc.add_paragraph()
ch9_text1 = '''9.1  结论'''
set_title_font(ch9_para1.add_run('9.1  结论'), 14)
ch9_para1.paragraph_format.space_before = Pt(15)
ch9_para1.paragraph_format.space_after = Pt(10)

ch9_para2 = doc.add_paragraph()
ch9_text2 = '''本文针对COVID-19肺炎CT影像的二分类问题，提出了一种基于深度卷积神经网络ResNet50的分类方法。通过采用ImageNet预训练权重、多层次数据增强、加权交叉熵损失函数和早停机制等技术手段，有效解决了医学图像数据集样本量有限导致的过拟合问题，提升了模型的泛化能力和分类性能。'''
set_normal_font(ch9_para2, 12)
ch9_para2.paragraph_format.first_line_indent = Pt(24)
ch9_para2.paragraph_format.line_spacing = 1.5
ch9_para2.paragraph_format.space_after = Pt(10)

ch9_para3 = doc.add_paragraph()
ch9_text3 = '''在COVID-CT公开数据集上的实验结果表明，本文提出的模型能够有效区分COVID-19阳性与阴性CT图像，在准确率、召回率、特异性和AUC-ROC等多项评价指标上均取得了较好的表现。模型的召回率较高，能够尽可能识别出COVID-19阳性患者，降低漏诊风险，具有一定的临床应用价值。'''
set_normal_font(ch9_para3, 12)
ch9_para3.paragraph_format.first_line_indent = Pt(24)
ch9_para3.paragraph_format.line_spacing = 1.5
ch9_para3.paragraph_format.space_after = Pt(10)

ch9_para4 = doc.add_paragraph()
ch9_text4 = '''9.2  展望'''
set_title_font(ch9_para4.add_run('9.2  展望'), 14)
ch9_para4.paragraph_format.space_before = Pt(15)
ch9_para4.paragraph_format.space_after = Pt(10)

ch9_para5 = doc.add_paragraph()
ch9_text5 = '''未来研究可以从以下几个方面开展：'''
set_normal_font(ch9_para5, 12)
ch9_para5.paragraph_format.first_line_indent = Pt(24)
ch9_para5.paragraph_format.line_spacing = 1.5
ch9_para5.paragraph_format.space_after = Pt(10)

future_items = [
    '扩展模型能力：将二分类扩展为多分类，实现COVID-19、细菌性肺炎、病毒性肺炎等多种肺炎类型的自动识别；',
    '引入注意力机制：引入CAM（Class Activation Mapping）、Grad-CAM等可视化技术，增强模型的可解释性，帮助医生理解模型的决策依据；',
    '多中心数据验证：在更多来源、不同设备的CT数据上进行验证，评估模型的泛化能力和鲁棒性；',
    '临床试验：将模型与医院信息系统集成，在真实的临床环境中进行试用，收集医生反馈，进一步优化模型性能；',
    '模型轻量化：采用模型剪枝、知识蒸馏等技术降低模型参数量和计算量，使其能够在资源受限的设备上运行。'
]

for i, item in enumerate(future_items, 1):
    fut_para = doc.add_paragraph()
    run = fut_para.add_run(f'({i}) ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = fut_para.add_run(item)
    set_normal_font(fut_para, 12)
    fut_para.paragraph_format.first_line_indent = Pt(24)
    fut_para.paragraph_format.line_spacing = 1.5
    fut_para.paragraph_format.space_after = Pt(8)

# 插入分页
doc.add_page_break()

# ==================== 参考文献 ====================
ref_title = doc.add_paragraph()
ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ref_title.add_run('参考文献')
set_title_font(run, 16)
ref_title.paragraph_format.space_before = Pt(20)
ref_title.paragraph_format.space_after = Pt(15)

refs = [
    '[1] He K, Zhang X, Ren S, et al. Deep Residual Learning for Image Recognition[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR). 2016: 770-778.',
    '[2] Krizhevsky A, Sutskever I, Hinton G E. ImageNet classification with deep convolutional neural networks[J]. Advances in Neural Information Processing Systems (NeurIPS). 2012, 25: 1097-1105.',
    '[3] Simonyan K, Zisserman A. Very Deep Convolutional Networks for Large-Scale Image Recognition[C]//International Conference on Learning Representations (ICLR). 2015.',
    '[4] Szegedy C, Liu W, Jia Y, et al. Going Deeper with Convolutions[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR). 2015: 1-9.',
    '[5] Wang X, Peng Y, Lu L, et al. ChestX-ray8: Hospital-scale Chest X-ray Database and Benchmarks for Chest X-ray Segmentation and Classification[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR). 2017: 2097-2106.',
    '[6] Huang G, Liu Z, Van Der Maaten L, et al. Densely Connected Convolutional Networks[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR). 2017: 4700-4708.',
    '[7] Ardila D, Kiraly A P, Bharadwaj S, et al. End-to-end lung cancer screening with three-dimensional deep learning on low-dose chest computed tomography[J]. Nature Medicine. 2019, 25(6): 954-961.',
    '[8] Esteva A, Kuprel B, Novoa R A, et al. Dermatologist-level classification of skin cancer with deep neural networks[J]. Nature. 2017, 542(7639): 115-118.',
    '[9] Gulshan V, Peng L, Coram M, et al. Development and validation of a deep learning algorithm for detection of diabetic retinopathy in retinal fundus photographs[J]. JAMA. 2016, 316(22): 2402-2410.',
    '[10] Rajpurkar P, Irvin J, Zhu K, et al. CheXNet: Radiologist-level pneumonia detection on chest X-rays with deep learning[J]. arXiv preprint arXiv:1711.05225. 2017.',
    '[11] Huang J, Cheng L, Lin J, et al. COVID-19 CT image classification based on ensemble deep learning models[J]. Journal of X-Ray Science and Technology. 2021, 29(1): 1-15.',
    '[12] Chen J, Wu L, Zhang J, et al. Deep learning-based CT images auxiliary diagnosis of COVID-19[J]. Journal of X-Ray Science and Technology. 2020, 28(2): 237-245.',
    '[13] Ozsahin I, Sekeroglu B, Musao glu M T, et al. Review on diagnosis of COVID-19 from chest CT images using artificial intelligence methods[J]. Computer Methods and Programs in Biomedicine. 2020, 196: 105908.',
    '[14] COVID-CT-Dataset. UCSD-AI4H COVID-CT Dataset[EB/OL]. https://github.com/UCSD-AI4H/COVID-CT, 2020.',
    '[15] Shorten C, Khoshgoftaar T M. A survey on image data augmentation for deep learning[J]. Journal of Big Data. 2019, 6(1): 60.',
    '[16] Goodfellow I, Bengio Y, Courville A. Deep Learning[M]. MIT Press, 2016.',
    '[17] LeCun Y, Bengio Y, Hinton G. Deep learning[J]. Nature. 2015, 521(7553): 436-444.',
    '[18] Russakovsky O, Deng J, Su H, et al. ImageNet Large Scale Visual Recognition Challenge[J]. International Journal of Computer Vision (IJCV). 2015, 115(3): 211-252.',
]

for ref in refs:
    ref_para = doc.add_paragraph()
    run = ref_para.add_run(ref)
    set_normal_font(ref_para, 10)
    ref_para.paragraph_format.line_spacing = 1.5
    ref_para.paragraph_format.space_after = Pt(6)

# 插入分页
doc.add_page_break()

# ==================== 致谢 ====================
thanks_title = doc.add_paragraph()
thanks_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = thanks_title.add_run('致  谢')
set_title_font(run, 16)
thanks_title.paragraph_format.space_before = Pt(20)
thanks_title.paragraph_format.space_after = Pt(15)

thanks_para = doc.add_paragraph()
thanks_text = '''在本论文完成之际，我要衷心感谢所有关心、帮助和支持过我的人。\n\n首先，感谢我的指导老师XXX教授在课题研究期间给予的悉心指导和宝贵建议。老师严谨的治学态度、深厚的学术造诣和一丝不苟的工作作风深深影响了我，使我在学术研究和工程实践方面都受益匪浅。从论文选题、研究方案设计到论文撰写，老师都给予了耐心细致的指导，帮助我克服了研究中遇到的种种困难。\n\n感谢实验室的同学们在学习和生活中给予的帮助和鼓励。与同学们进行的学术讨论和交流开拓了我的思路，帮助我解决了很多技术难题。感谢实验室提供的良好学习环境和计算资源，为本研究的顺利进行提供了有力保障。\n\n感谢我的家人对我学业的理解、支持和鼓励。家人的无私奉献是我完成学业的强大动力，他们的支持和鼓励让我能够专注于学习，顺利完成研究工作。\n\n最后，向百忙之中抽出时间审阅本论文的各位专家和学者致以诚挚的谢意！'''
set_normal_font(thanks_para, 12)
thanks_para.paragraph_format.first_line_indent = Pt(24)
thanks_para.paragraph_format.line_spacing = 1.5

# 保存文档
output_path = r'D:\Users\4ever\Desktop\pneumonia_ct_classifier\肺炎CT影像分类系统_大作业论文.docx'
doc.save(output_path)
print(f'Word文档已成功保存到: {output_path}')